from crewai import Agent, Crew, Process, Task
from crewai.project import CrewBase, agent, crew, task
from crewai.tools import BaseTool
from pydantic import BaseModel, Field
from typing import Type, List, Any
import os
from crewai_tools import XMLSearchTool, DOCXSearchTool
import logging
import pandas as pd
from docx import Document
from docx.shared import Pt
import json
from typing import Any, Dict

# Custom Directory Reader Tool
class CustomDirectoryReaderInput(BaseModel):
    directory: str = Field(..., description="Directory to scan.")
    ignored_folders: list[str] = Field(default_factory=list, description="Folders to ignore.")
    ignored_files: list[str] = Field(default_factory=list, description="Files to ignore.")
    

class CustomDirectoryReaderTool(BaseTool):
    name: str = "Custom Directory Reader Tool"
    description: str = "Recursively reads a directory while ignoring specified folders and files."
    args_schema: Type[BaseModel] = CustomDirectoryReaderInput

    def _run(self, directory: str, ignored_folders: list[str], ignored_files: list[str]) -> list[str]:
        """
        Recursively reads a directory while ignoring specified folders and files.
        Filters files to include only those with extensions .xaml, .json, and .xlsx.
        """
        logging.basicConfig(level=logging.DEBUG)
        logging.debug(f"Received directory: {directory}")
        logging.debug(f"Ignored folders: {ignored_folders}")
        logging.debug(f"Ignored files: {ignored_files}")

        # Validate if the directory exists and is accessible
        if not os.path.isdir(directory):
            logging.error(f"Directory '{directory}' does not exist or is not accessible.")
            raise ValueError(f"Directory '{directory}' does not exist or is not accessible.")

        try:
            files_list = []
            for root, dirs, files in os.walk(directory):
                # Remove ignored folders that exist
                dirs[:] = [d for d in dirs if not (d in ignored_folders and os.path.isdir(os.path.join(root, d)))]

                # Add files to the list, ignoring those in ignored_files that exist
                for file in files:
                    if not (file in ignored_files and os.path.isfile(os.path.join(root, file))):
                        if file.endswith(('.xaml', '.json', '.xlsx')):
                            files_list.append(os.path.join(root, file))

            logging.debug(f"Files found: {files_list}")
            return files_list

        except Exception as e:
            logging.error(f"Error reading directory '{directory}': {e}")
            return {"error": str(e)}

# Custom File Reader Tool
class CustomFileReaderInput(BaseModel):
    file_paths: List[str] = Field(..., description="List of file paths to read.")

class CustomFileReaderTool(BaseTool):
    name: str = "Custom File Reader Tool"
    description: str = "Reads the content of specified files and returns their content."
    args_schema: Type[BaseModel] = CustomFileReaderInput

    def _run(self, file_paths: List[str]) -> dict:
        """
        Reads the content of specified files. Supports .xaml, .json, and .xlsx files.

        Args:
            file_paths (List[str]): List of file paths to read.

        Returns:
            dict: Dictionary with file paths as keys and their content as values.
        """
        file_contents = {}
        for file_path in file_paths:
            try:
                if file_path.endswith('.xlsx'):
                    # Use pandas to read Excel files
                    df = pd.read_excel(file_path)
                    file_contents[file_path] = df.to_dict()  # Convert DataFrame to dictionary
                else:
                    # Default to text reading for other file types
                    with open(file_path, "r", encoding="utf-8") as file:
                        file_contents[file_path] = file.read()
            except FileNotFoundError:
                logging.error(f"File not found: {file_path}")
                file_contents[file_path] = "Error: File not found."
            except PermissionError:
                logging.error(f"Permission denied: {file_path}")
                file_contents[file_path] = "Error: Permission denied."
            except Exception as e:
                logging.error(f"Error reading file {file_path}: {e}")
                file_contents[file_path] = f"Error: {str(e)}"
        return file_contents

# Modelo de entrada para validação
class DocumentWriterInput(BaseModel):
    content: Any = Field(..., description="Content to write into the document.")
    output_path: str = Field(..., description="Path where the document will be saved.")
    styles: Dict[str, Dict[str, Any]] = Field(default_factory=dict, description="Custom styles for the document.")

class DocumentWriterTool(BaseTool):
    name: str = "Document Writer Tool"
    description: str = "Writes structured content into a Word document with custom formatting."
    args_schema: Type[BaseModel] = DocumentWriterInput

    def _run(self, content: dict, output_path: str, styles: dict = None) -> dict:
        try:
            # Validar ou converter estilos
            if isinstance(styles, str):
                styles = json.loads(styles)

            # Criar o documento
            document = Document()

            # Adicionar conteúdo com sessões
            for section_title, section_content in content.items():
                # Adicionar título da sessão
                section_heading = document.add_heading(section_title, level=1)
                if "section" in styles:
                    font = section_heading.style.font
                    font.size = Pt(styles["section"].get("font_size", 16))
                    font.bold = styles["section"].get("bold", True)

                # Adicionar conteúdo da sessão
                if isinstance(section_content, list):
                    for paragraph_text in section_content:
                        paragraph = document.add_paragraph(paragraph_text)
                        if "paragraph" in styles:
                            font = paragraph.style.font
                            font.size = Pt(styles["paragraph"].get("font_size", 12))
                else:
                    paragraph = document.add_paragraph(section_content)
                    if "paragraph" in styles:
                        font = paragraph.style.font
                        font.size = Pt(styles["paragraph"].get("font_size", 12))

            # Salvar o documento
            document.save(output_path)
            return {"success": True, "output_path": output_path}
        except Exception as e:
            logging.error(f"[DocumentWriterTool] Error writing document: {e}")
            return {"error": str(e)}

@CrewBase
class Uipathdocumantationgenerator():
    """Uipathdocumantationgenerator crew"""

    agents_config = 'config/agents.yaml'
    tasks_config = 'config/tasks.yaml'

    def __init__(self, input_directory=None, ignored_folders=None, ignored_files=None):
        if input_directory and not os.path.isdir(input_directory):
            raise ValueError(f"Invalid directory: {input_directory}")
        self.input_directory = input_directory
        self.ignored_folders = ignored_folders or []
        self.ignored_files = ignored_files or []
        self.task_outputs = {}  # Dictionary to store task outputs

    @agent
    def document_reader(self) -> Agent:
        print(f"Directory provided to agent: {self.input_directory}")
        return Agent(
            config=self.agents_config['document_reader'],
            verbose=True,
            tools=[
                CustomDirectoryReaderTool(),
                CustomFileReaderTool(),
                XMLSearchTool(),
            ],
        )


    @agent
    def documentation_creator(self) -> Agent:
        return Agent(
            config=self.agents_config['documentation_creator'],
            verbose=True,
            tools=[
                DOCXSearchTool(),
                DocumentWriterTool(),
                
            ]
        )

    @task
    def directory_scanning_task(self) -> Task:
        return Task(
            config=self.tasks_config['directory_scanning_task'],
            tools=[
                CustomDirectoryReaderTool(),
            ],
        )

    @task
    def file_analysis_task(self) -> Task:
        return Task(
            config=self.tasks_config['file_analysis_task'],
            tools=[
                CustomFileReaderTool(),
                XMLSearchTool(),
            ],
            output_file="SDD.md"
    
        )   
        

    @task
    def documentation_writing_task(self) -> Task:
        return Task(
            config=self.tasks_config['documentation_writing_task'],
            tools=[
                DocumentWriterTool(),
            ],
            
        )

    @crew
    def crew(self) -> Crew:
        """Creates the Uipathdocumantationgenerator crew"""
        logging.debug(f"Initializing crew with directory: {self.input_directory}")
        logging.debug(f"Ignored folders: {self.ignored_folders}")
        logging.debug(f"Ignored files: {self.ignored_files}")

        return Crew(
            agents=self.agents, 
            tasks=self.tasks, 
            process=Process.sequential,
            memory=True,
            verbose=True,
        )